"""
Will's Resume Generator
Generates formatted .docx resume files from plain text resume data.
Run this script to create all tailored resume versions.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "Resumes")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def add_contact_header(doc, name, contact_line):
    """Add the name and contact info as a centered header."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(contact_line)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(80, 80, 80)


def add_section_heading(doc, text):
    """Add a section heading with a bottom border."""
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)
    # Add bottom border via paragraph formatting
    pf = p.paragraph_format
    pf.border_bottom = True


def add_body_text(doc, text, bold=False, italic=False, indent=False):
    """Add a paragraph of body text."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    run.italic = italic


def add_bullet(doc, text, indent_level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + (indent_level * 0.25))
    p.space_after = Pt(1)
    run = p.clear().add_run(text)
    run.font.size = Pt(10)


def add_job_header(doc, title, company, dates):
    """Add a job title line: Title | Company | Dates."""
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(2)

    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(10)

    run_sep = p.add_run(" | ")
    run_sep.font.size = Pt(10)

    run_company = p.add_run(company)
    run_company.font.size = Pt(10)

    run_sep2 = p.add_run(" | ")
    run_sep2.font.size = Pt(10)

    run_dates = p.add_run(dates)
    run_dates.font.size = Pt(10)
    run_dates.font.color.rgb = RGBColor(100, 100, 100)


def parse_and_build_docx(resume_text, version_id):
    """Parse resume text and generate a formatted .docx file."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    lines = resume_text.strip().split('\n')
    i = 0

    # Parse the name (first line)
    if lines:
        name = lines[0].strip()
        i += 1

    # Parse contact line (second line)
    contact = ""
    if i < len(lines):
        contact = lines[i].strip()
        i += 1

    add_contact_header(doc, name, contact)

    # Parse the rest
    current_section = None
    while i < len(lines):
        line = lines[i].strip()
        i += 1

        if not line:
            continue

        # Detect section headings (ALL CAPS lines or known headers)
        known_sections = [
            'KEY QUALIFICATIONS', 'PROFESSIONAL SKILLS', 'AI & PROFESSIONAL DEVELOPMENT',
            'PROFESSIONAL EXPERIENCE', 'MILITARY EXPERIENCE', 'EDUCATION & CREDENTIALS',
            'AI & EMERGING TECHNOLOGY', 'COMPLIANCE & OPERATIONS', 'ENABLEMENT & CHANGE MANAGEMENT',
            'LEADERSHIP', 'PROGRAM & OPERATIONS MANAGEMENT', 'DATA & ANALYTICS',
            'AI & AUTOMATION', 'BUSINESS OPERATIONS', 'PROJECT & DATA MANAGEMENT',
            'STRATEGY & OPERATIONS', 'PROGRAM BUILDING', 'OPERATIONS & AUTOMATION',
            'PROJECT & DATA MANAGEMENT'
        ]

        is_section = (line.isupper() and len(line) > 3) or line in known_sections

        if is_section:
            current_section = line
            add_section_heading(doc, line)
            continue

        # Detect job headers (Title | Company | Dates pattern)
        job_match = re.match(r'^(.+?)\s*\|\s*(.+?)\s*\|\s*(\d{4}.*)', line)
        if job_match:
            add_job_header(doc, job_match.group(1), job_match.group(2), job_match.group(3))
            continue

        # Detect sub-section headers with dates (e.g., "2023 - Present | Self-Directed...")
        date_header = re.match(r'^(\d{4}\s*-\s*\w+)\s*\|\s*(.+)', line)
        if date_header:
            p = doc.add_paragraph()
            p.space_before = Pt(6)
            run = p.add_run(f"{date_header.group(1)} | {date_header.group(2)}")
            run.bold = True
            run.italic = True
            run.font.size = Pt(10)
            continue

        # Detect bullet points
        if line.startswith('- '):
            add_bullet(doc, line[2:])
            continue

        # Detect skill categories (Category: skill1, skill2, ...)
        skill_match = re.match(r'^(.+?):\s*(.+)', line)
        if skill_match and current_section and 'SKILLS' in current_section:
            p = doc.add_paragraph()
            p.space_after = Pt(2)
            run_cat = p.add_run(skill_match.group(1) + ": ")
            run_cat.bold = True
            run_cat.font.size = Pt(10)
            run_skills = p.add_run(skill_match.group(2))
            run_skills.font.size = Pt(10)
            continue

        # Default: add as body text
        add_body_text(doc, line)

    # Save
    filepath = os.path.join(OUTPUT_DIR, f"{version_id}.docx")
    doc.save(filepath)
    return filepath


# ============================================================
# RESUME DATA - All 4 tailored versions
# ============================================================

RESUMES = {
    "WR-CIRCLE-COMPOPS-20260304": """WILL SHARP
Atlanta, GA | 678.997.4765 | wilsharp20@gmail.com | linkedin.com/in/williampsharp

Compliance operations leader with 10+ years driving process automation, stakeholder enablement, and operational risk reduction across financial services. Reduced $200K in annual operating costs through automation at Hiscox USA and led a 500K-account data migration at 96% accuracy. Now building on that foundation with focused AI and machine learning training to bring intelligent automation into compliance and operations workflows. U.S. Marine Corps veteran with secret clearance experience.

KEY QUALIFICATIONS
- 10+ years leading operations process improvement and compliance-adjacent workflows in financial services
- Proven track record designing and deploying automation solutions that reduce cost and risk
- AI and machine learning training focused on responsible implementation in business operations
- Experience building enablement programs, training materials, and change management infrastructure from the ground up
- Cross-functional stakeholder management across underwriting, claims, IT, and executive teams

PROFESSIONAL SKILLS
Compliance & Operations: Process Improvement, Compliance Operations, Robotic Process Automation (RPA), Risk Reduction, Cost-Benefit Analysis
Enablement & Change Management: Training Program Development, Change Management, User Adoption, Data Governance, Requirements Gathering
AI & Emerging Technology: AI for Business Operations, Responsible AI, Machine Learning Fundamentals, AI Tool Integration
Leadership: Stakeholder Management, Cross-Departmental Coordination, Team Development & Mentorship, Business Analysis

AI & PROFESSIONAL DEVELOPMENT
2023 - Present | Self-Directed Career Investment
- Career Essentials in Business Analysis - Microsoft
- AI Accountability: Build Responsible and Transparent Systems - LinkedIn
- AI Challenges and Opportunities for Leadership - LinkedIn
- Artificial Intelligence Foundations: Neural Networks - LinkedIn
- Introduction to Artificial Intelligence - LinkedIn
- Project Management Foundations: Requirements - LinkedIn

PROFESSIONAL EXPERIENCE
Lead Automation Consultant | Hiscox USA | 2021-2023
- Reduced $200K in annual operating costs by identifying process pain points, designing automation solutions, and managing delivery across business units
- Conducted feasibility assessments and cost-benefit analyses to evaluate and prioritize automation opportunities
- Gathered functional requirements and mapped existing processes to define targeted automation solutions for cross-functional teams
- Maintained health of automated processes through error investigation, outage resolution, and development of testing protocols

Data Conversion Project Lead | Hiscox USA | 2018-2021
- Led cross-functional data migration teams through a full platform transition, safeguarding sensitive client information (PII) in compliance with regulatory requirements
- Built change management infrastructure including user guides, training materials, data dictionaries, and governance policies for end-user adoption
- Exceeded migration targets by 4%, achieving a 96% successful conversion rate across 500K customer accounts

Senior Underwriting Assistant | Hiscox USA | 2014-2017
- Created the Underwriting Assistant function from the ground up, building training programs, career development paths, and mentorship structures
- Implemented workflow and workload management improvements including triage systems and cross-departmental coordination

Specialty Claims Adjuster | AIG | 2012-2014
- Managed end-to-end claims process ensuring compliance with policy terms and applicable regulations across all 50 states
- Coordinated indemnity, subrogation, and litigation with external parties

MILITARY EXPERIENCE
Team Leader / Armory Manager | United States Marine Corps | 2005-2009
- Led a team of Marines in day-to-day operations; held U.S. Secret clearance
- Managed accountability and distribution of company weaponry totaling $7M+ as Armory Manager

EDUCATION & CREDENTIALS
B.B.A. in Risk Management and Insurance - Georgia State University
Associate in Insurance (AINS) - American Institute for CPCU""",

    "WR-FHF-CREDPM-20260304": """WILL SHARP
Atlanta, GA | 678.997.4765 | wilsharp20@gmail.com | linkedin.com/in/williampsharp

Operations and program management leader with 10+ years driving process automation, data-driven decision making, and cost reduction across financial services. Reduced $200K in annual operating costs through automation at Hiscox USA and led a 500K-account data migration at 96% accuracy. Now building on that foundation with focused AI and machine learning training to bring intelligent automation into program management and financial operations workflows. U.S. Marine Corps veteran with secret clearance experience.

KEY QUALIFICATIONS
- 10+ years in financial services operations with deep experience in program design, implementation, and continuous improvement
- Led automation initiatives that delivered $200K in annual cost savings through process redesign
- Managed cross-functional teams through complex data migration affecting 500K+ customer accounts
- AI and machine learning training focused on practical automation for financial services operations
- Built training programs, governance policies, and change management infrastructure from scratch

PROFESSIONAL SKILLS
Program & Operations Management: Process Improvement, Program Design & Optimization, Cost-Benefit Analysis, Operations Consulting, Requirements Gathering
Data & Analytics: Data Migration, Data Governance, Performance Metrics, Feasibility Assessment
AI & Automation: AI for Business Operations, Robotic Process Automation (RPA), Machine Learning Fundamentals, Responsible AI
Leadership: Cross-Functional Partnership, Stakeholder Management, Team Development & Mentorship, Change Management

AI & PROFESSIONAL DEVELOPMENT
2023 - Present | Self-Directed Career Investment
- Career Essentials in Business Analysis - Microsoft
- AI Challenges and Opportunities for Leadership - LinkedIn
- AI Accountability: Build Responsible and Transparent Systems - LinkedIn
- Artificial Intelligence Foundations: Neural Networks - LinkedIn
- Introduction to Artificial Intelligence - LinkedIn
- Project Management Foundations: Requirements - LinkedIn

PROFESSIONAL EXPERIENCE
Lead Automation Consultant | Hiscox USA | 2021-2023
- Reduced $200K in annual operating costs by identifying process pain points, designing automation solutions, managing delivery, and resolving post-implementation issues with stakeholders
- Conducted feasibility assessments and cost-benefit analyses to evaluate and prioritize automation opportunities across business units
- Gathered functional requirements and mapped existing processes and datasets to define targeted automation solutions
- Maintained health of automated processes through error investigation, outage resolution, troubleshooting, and development of testing protocols

Data Conversion Project Lead | Hiscox USA | 2018-2021
- Led cross-functional teams through a full platform transition from legacy systems to a new administration system, safeguarding sensitive client information (PII)
- Exceeded migration targets by 4%, achieving a 96% successful conversion rate across 500K customer accounts
- Built change management infrastructure including user guides, training materials, data dictionaries, and governance policies for end-user adoption

Senior Underwriting Assistant | Hiscox USA | 2014-2017
- Created the Underwriting Assistant function from the ground up, building training programs, career development paths, and mentorship structures for junior associates
- Implemented workflow and workload management improvements across operations, including triage systems and cross-departmental coordination

Specialty Claims Adjuster | AIG | 2012-2014
- Managed end-to-end claims process from intake through resolution, coordinating with external parties across all 50 states
- Licensed in all 50 states to investigate and resolve liability claims in compliance with policy terms and regulations

MILITARY EXPERIENCE
Team Leader / Armory Manager | United States Marine Corps | 2005-2009
- Led a team of Marines in day-to-day operations; held U.S. Secret clearance
- Managed accountability and distribution of company weaponry totaling $7M+; trained and mentored a team of 4

EDUCATION & CREDENTIALS
B.B.A. in Risk Management and Insurance - Georgia State University
Associate in Insurance (AINS) - American Institute for CPCU""",

    "WR-SAMS-BIZOPS-20260304": """WILL SHARP
Atlanta, GA | 678.997.4765 | wilsharp20@gmail.com | linkedin.com/in/williampsharp

Business operations leader with 10+ years driving process automation, cross-functional operational efficiency, and data-driven performance improvement. Reduced $200K in annual operating costs through automation at Hiscox USA and led a 500K-account data migration at 96% accuracy. Now building on that foundation with focused AI and machine learning training to bring intelligent automation into business operations and GTM workflows. U.S. Marine Corps veteran with secret clearance experience.

KEY QUALIFICATIONS
- 10+ years in operations consulting and business operations with proven track record of delivering measurable efficiency gains
- Led cross-functional automation initiatives saving $200K annually through data-driven process redesign
- Experience partnering with leadership to evaluate opportunities, diagnose performance, and develop action plans
- AI and machine learning training focused on operational efficiency and performance optimization
- Strong stakeholder management across executive, technical, and operational teams

PROFESSIONAL SKILLS
Business Operations: Operations Consulting, Process Improvement, Cost-Benefit Analysis, Performance Analysis, Feasibility Assessment
Project & Data Management: Project Management, Data Migration, Requirements Gathering, Change Management, Executive Reporting
AI & Emerging Technology: AI for Business Operations, Robotic Process Automation (RPA), Machine Learning Fundamentals, Neural Networks
Leadership: Stakeholder Management, Cross-Departmental Coordination, Team Development & Mentorship, Business Analysis

AI & PROFESSIONAL DEVELOPMENT
2023 - Present | Self-Directed Career Investment
- Career Essentials in Business Analysis - Microsoft
- AI Challenges and Opportunities for Leadership - LinkedIn
- AI Accountability: Build Responsible and Transparent Systems - LinkedIn
- Artificial Intelligence Foundations: Neural Networks - LinkedIn
- Introduction to Artificial Intelligence - LinkedIn
- Project Management Foundations: Requirements - LinkedIn

PROFESSIONAL EXPERIENCE
Lead Automation Consultant | Hiscox USA | 2021-2023
- Reduced $200K in annual operating costs by identifying process pain points, designing automation solutions, managing delivery, and resolving post-implementation issues with stakeholders
- Conducted feasibility assessments and cost-benefit analyses to evaluate and prioritize automation opportunities across business units
- Gathered functional requirements and mapped existing processes and datasets to define targeted automation solutions
- Developed executive materials and presented findings to leadership for strategic decision-making

Data Conversion Project Lead | Hiscox USA | 2018-2021
- Led cross-functional data migration teams through a full platform transition, safeguarding sensitive client information (PII)
- Exceeded migration targets by 4%, achieving a 96% successful conversion rate across 500K customer accounts
- Built change management infrastructure including user guides, training materials, data dictionaries, and governance policies

Senior Underwriting Assistant | Hiscox USA | 2014-2017
- Created the Underwriting Assistant function from the ground up, building training programs, career development paths, and mentorship structures
- Implemented workflow and workload management improvements including triage systems and cross-departmental coordination

Specialty Claims Adjuster | AIG | 2012-2014
- Managed end-to-end process from intake through resolution, coordinating with external parties
- Licensed in all 50 states; investigated and resolved cases in compliance with regulations

MILITARY EXPERIENCE
Team Leader / Armory Manager | United States Marine Corps | 2005-2009
- Led a team of Marines in day-to-day operations; held U.S. Secret clearance
- Managed accountability and distribution of assets totaling $7M+; trained and mentored a team of 4

EDUCATION & CREDENTIALS
B.B.A. in Risk Management and Insurance - Georgia State University
Associate in Insurance (AINS) - American Institute for CPCU""",

    "WR-SPARK-STRATOPS-20260304": """WILL SHARP
Atlanta, GA | 678.997.4765 | wilsharp20@gmail.com | linkedin.com/in/williampsharp

Strategy and operations leader with 10+ years building operational infrastructure, driving process automation, and scaling programs in financial services. Reduced $200K in annual operating costs through automation at Hiscox USA and led a 500K-account data migration at 96% accuracy. Now building on that foundation with focused AI and machine learning training to bring intelligent automation into high-growth operational environments. U.S. Marine Corps veteran with secret clearance experience.

KEY QUALIFICATIONS
- 10+ years building and scaling operations functions from the ground up in fast-paced environments
- Created entire operational functions (Underwriting Assistant team) including hiring, training, and career development infrastructure
- Led automation and process improvement initiatives delivering $200K in annual cost savings
- AI and machine learning training focused on bringing intelligent automation to operational workflows
- Proven cross-functional leader comfortable working with executive teams, product, and field operations

PROFESSIONAL SKILLS
Strategy & Operations: Process Improvement, Operations Consulting, Strategic Planning, Cost-Benefit Analysis, Feasibility Assessment
Program Building: Change Management, Training Program Development, Requirements Gathering, Data Governance, User Adoption
AI & Automation: AI for Business Operations, Robotic Process Automation (RPA), Machine Learning Fundamentals, Responsible AI
Leadership: Cross-Functional Coordination, Stakeholder Management, Team Development & Mentorship, Business Analysis

AI & PROFESSIONAL DEVELOPMENT
2023 - Present | Self-Directed Career Investment
- Career Essentials in Business Analysis - Microsoft
- AI Challenges and Opportunities for Leadership - LinkedIn
- AI Accountability: Build Responsible and Transparent Systems - LinkedIn
- Artificial Intelligence Foundations: Neural Networks - LinkedIn
- Introduction to Artificial Intelligence - LinkedIn
- Project Management Foundations: Requirements - LinkedIn

PROFESSIONAL EXPERIENCE
Lead Automation Consultant | Hiscox USA | 2021-2023
- Reduced $200K in annual operating costs by identifying process pain points, designing automation solutions, managing delivery, and resolving post-implementation issues
- Conducted feasibility assessments and cost-benefit analyses to evaluate and prioritize initiatives across business units
- Gathered functional requirements and mapped existing processes to define targeted solutions
- Maintained health of automated processes through error investigation, outage resolution, and development of testing protocols

Data Conversion Project Lead | Hiscox USA | 2018-2021
- Led cross-functional teams through a full platform transition from legacy systems, safeguarding sensitive client information
- Exceeded migration targets by 4%, achieving a 96% successful conversion rate across 500K accounts
- Built change management infrastructure including user guides, training materials, data dictionaries, and governance policies for organization-wide adoption

Senior Underwriting Assistant | Hiscox USA | 2014-2017
- Created an entirely new operational function from the ground up, building hiring processes, training programs, career development paths, and mentorship structures
- Implemented workflow and workload management improvements including triage systems and cross-departmental coordination protocols

Specialty Claims Adjuster | AIG | 2012-2014
- Managed end-to-end operational process from intake through resolution, coordinating across internal and external stakeholders
- Licensed in all 50 states; investigated and resolved cases in compliance with regulations

MILITARY EXPERIENCE
Team Leader / Armory Manager | United States Marine Corps | 2005-2009
- Led a team of Marines solving rapidly changing operational challenges; held U.S. Secret clearance
- Managed accountability and distribution of $7M+ in assets; trained and mentored a team of 4

EDUCATION & CREDENTIALS
B.B.A. in Risk Management and Insurance - Georgia State University
Associate in Insurance (AINS) - American Institute for CPCU""",
}


# Also generate the base resume
RESUMES["WR-BASE-CURRENT"] = """WILL SHARP
Atlanta, GA | 678.997.4765 | wilsharp20@gmail.com | linkedin.com/in/williampsharp

Operations leader with 10+ years driving process automation, data migration, and cost reduction across insurance and financial services. Reduced $200K in annual operating costs through automation at Hiscox USA and led a 500K-account data migration at 96% accuracy. Now building on that foundation with focused AI and machine learning training to bring intelligent automation into operations workflows. U.S. Marine Corps veteran with secret clearance experience.

PROFESSIONAL SKILLS
Operations & Automation: Process Improvement, Robotic Process Automation (RPA), Operations Consulting, Cost-Benefit Analysis
Project & Data Management: Project Management, Data Migration, Requirements Gathering, Change Management
AI & Emerging Technology: AI for Business Operations, Responsible AI, Machine Learning Fundamentals, Neural Networks
Leadership: Stakeholder Management, Team Development & Mentorship, Cross-Departmental Coordination, Business Analysis

AI & PROFESSIONAL DEVELOPMENT
2023 - Present | Self-Directed Career Investment
- Career Essentials in Business Analysis - Microsoft
- AI Challenges and Opportunities for Leadership - LinkedIn
- AI Accountability: Build Responsible and Transparent Systems - LinkedIn
- Artificial Intelligence Foundations: Neural Networks - LinkedIn
- Introduction to Artificial Intelligence - LinkedIn
- Project Management Foundations: Requirements - LinkedIn

PROFESSIONAL EXPERIENCE
Lead Automation Consultant | Hiscox USA | 2021-2023
- Reduced $200K in annual operating costs by identifying process pain points, designing automation solutions, managing delivery, and resolving post-implementation issues with stakeholders
- Conducted feasibility assessments and cost-benefit analyses to evaluate and prioritize automation opportunities across business units
- Gathered functional requirements and mapped existing processes and datasets to define targeted automation solutions
- Maintained health of automated processes through error investigation, outage resolution, troubleshooting, and development of testing protocols

Data Conversion Project Lead | Hiscox USA | 2018-2021
- Led cross-functional data migration teams through a full platform transition from legacy systems to a new policy administration system, safeguarding sensitive client information (PII)
- Exceeded migration targets by 4%, achieving a 96% successful conversion rate across 500K customer accounts
- Built change management infrastructure including user guides, training materials, data dictionaries, and governance policies for end-user adoption

Senior Underwriting Assistant | Hiscox USA | 2014-2017
- Created the Underwriting Assistant function from the ground up, building training programs, career development paths, and mentorship structures for junior associates
- Implemented workflow and workload management improvements across underwriting operations, including triage systems and cross-departmental coordination

Specialty Claims Adjuster | AIG | 2012-2014
- Managed end-to-end claims process from intake through resolution, coordinating indemnity, subrogation, and litigation with external parties
- Licensed in all 50 states to investigate and resolve liability claims in compliance with policy terms and applicable regulations

MILITARY EXPERIENCE
Team Leader / Armory Manager | United States Marine Corps | 2005-2009
- Led a team of Marines in day-to-day operations solving rapidly changing security and stabilization challenges; held U.S. Secret clearance
- Managed accountability and distribution of company weaponry totaling $7M+ as Armory Manager; trained and mentored a team of 4

EDUCATION & CREDENTIALS
B.B.A. in Risk Management and Insurance - Georgia State University
Associate in Insurance (AINS) - American Institute for Chartered Property Casualty Underwriters"""


if __name__ == "__main__":
    print(f"Generating {len(RESUMES)} resume files...\n")
    for version_id, text in RESUMES.items():
        path = parse_and_build_docx(text, version_id)
        print(f"  Created: {path}")
    print(f"\nDone. All files saved to: {OUTPUT_DIR}")
